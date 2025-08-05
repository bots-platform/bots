# word_extractor.py

import pandas as pd
from docx import Document
from typing import List, Dict

from app.modules.sga.minpub.report_validator.service.objetivos.utils.decorators import ( 
    log_exceptions
)

@log_exceptions
def extract_averias_table(path_docx: str) -> pd.DataFrame:
    """
    Load the specified table from a .docx and return it as a DataFrame.
    index = 4 (CLARO)
    index = 5 (CLIENTE)
    index =  (TERCEROS)
    We assume the first row is the header.
    """

    try:
        table_index_list = [4,5,6]

        doc = Document(path_docx)
        all_dfs: List[pd.DataFrame] = []
       
        for table_index in table_index_list:

            if table_index >= len(doc.tables):
                print(f"Document only has {len(doc.tables)} tables. Skipping table {table_index}.")
                continue
            
            table = doc.tables[table_index]

            headers = [cell.text.strip() for cell in table.rows[0].cells]

            records: List[Dict] = []
            for row in table.rows[1:]:
                vals = [cell.text.strip() for cell in row.cells]

                if len(vals) < len(headers):
                    vals += [""] * (len(headers) - len(vals))
                records.append(dict(zip(headers, vals)))

            df_part = pd.DataFrame.from_records(records)
            
            if table_index == 4:
                df_part['responsable'] = 'CLARO'
            elif table_index == 5:
                df_part['responsable'] = 'CLIENTE'
            else: 
                df_part['responsable'] = 'TERCEROS'

            all_dfs.append(df_part)

        if not all_dfs:
            # Si no hay tablas válidas, crear DataFrame vacío
            return pd.DataFrame(columns=['responsable'])
        
        result = pd.concat(all_dfs, ignore_index=True)
        return result
        
    except Exception as e:
        print(f"Error al procesar archivo DOCX de averías: {e}")
        return pd.DataFrame(columns=['responsable'])
