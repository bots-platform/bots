
from typing import List, Dict
import pandas as pd
import re
from datetime import datetime, timedelta
from docx import Document
from typing import Tuple, Optional

# import language_tool_python
# _tool_es = language_tool_python.LanguageTool('es')


from app.modules.sga.minpub.report_validator.service.objetivos.decorators import ( 
    log_exceptions
)
from utils.logger_config import get_sga_logger
logger = get_sga_logger()



#PARADAS DE RELOJ
@log_exceptions
def resolve_clock_stop_overlaps(clock_stops: List[Dict]) -> List[Dict]:
    """
    Eliminate overlaps in clock stops (paradas de reloj) by nro_incidencia.

    Args:
        clock_stops: List of clock stops with 'start' 'end' datetime and 'nro_incidencia'

    Returns:
        List of non-overlapping clock stops
            
    """
    if not clock_stops:
        return []
    
    incidents = {}
    for stop in clock_stops:
        nro_incidencia = stop.get('nro_incidencia', 'unknown')
        if nro_incidencia not in incidents:
            incidents[nro_incidencia] = []
        incidents[nro_incidencia].append(stop)

    
    resolved_all = []   

    for nro_incidencia, incident_stops in incidents.items():
        sorted_stops = sorted(incident_stops, key=lambda x: x['start'])

        for i, stop in enumerate(sorted_stops):
            if pd.isna(stop['end']):
                if i < len(sorted_stops) - 1 and not pd.isna(sorted_stops[i+1]['start']):
                    stop['end'] = sorted_stops[i+1]['start']
                else:
                    logger.warning(f"Removing stop with missing end date for nro_incidencia {nro_incidencia}")
                    continue
        
        valid_stops = [stop for stop in sorted_stops if not pd.isna(stop['end'])]

        if not valid_stops:
            continue

        resolved_stops = [valid_stops[0]]

        for current_stop in valid_stops[1:]:
            last_resolved = resolved_stops[-1]

            if current_stop['start'] <= last_resolved['end']:
                last_resolved['end'] = max(last_resolved['end'], current_stop['end'])
            else:
                resolved_stops.append(current_stop)

        resolved_all.extend(resolved_stops)

    return resolved_all

@log_exceptions
def calculate_total_clock_stop_minutes_by_incidencia(
    nro_incidencia:str,
    interruption_start: datetime,
    interruption_end: datetime,
    df_sga_paradas: pd.DataFrame
    ) -> float:
    """
    Calculate the total clock minutes for a ticket, considering constraints.

    Args:
        nro_incidencia: The ticket identifier
        interrupcion_inicio: Start time of the interruption from REPORTE DINAMICO 335 
        interrupcion_fin: End time of the interruption from REPORTE DINAMICO 335 
    
    Returns:
        Total clock stop minutes
    
    """   
    df_sga_paradas['nro_incidencia'] = df_sga_paradas['nro_incidencia'].astype(str)
    nro_incidencia_stops = df_sga_paradas[df_sga_paradas['nro_incidencia'] == nro_incidencia].copy()

    if nro_incidencia_stops.empty:
        logger.info(f"No clock stops found for incident {nro_incidencia}")
        return 0.0
    
    clock_stops = []

    for _, stop in nro_incidencia_stops.iterrows():
        start_date = stop.get('startdate')
        end_date = stop.get('enddate')

        if pd.isna(start_date):
            logger.warning(f"Skipping record with missing start date for incident {nro_incidencia}")
            continue

        if start_date < interruption_start:
            logger.info(f"Adjusting start time to interruption en for incident {nro_incidencia}")
            start_date = interruption_start

        if not pd.isna(end_date):
            if end_date > interruption_end:
                logger.info(f"Adjusting end time to interruption en for incident {nro_incidencia}")
                end_date = interruption_end

            if start_date < end_date:
                clock_stops.append({
                    'start': start_date,
                    'end': end_date,
                    'nro_incidencia': nro_incidencia
                })
        else:
            clock_stops.append({
                'start': start_date,
                'end': end_date,
                'nro_incidencia': nro_incidencia
            })
    resolved_stops = resolve_clock_stop_overlaps(clock_stops)

    total_minutes = sum(
        (stop['end'] - stop['start']).total_seconds() / 60
        for stop in resolved_stops
        if not pd.isna(stop['end']) and not pd.isna(stop['start'])
    )
    return total_minutes

@log_exceptions
def calculate_total_clock_stop_by_incidencia(
    nro_incidencia:str,
    interruption_start: datetime,
    interruption_end: datetime,
    df_sga_paradas: pd.DataFrame
    ) -> float:
    """
    Calculate the total clock stops for a ticket, considering constraints.

    Args:
        nro_incidencia: The ticket identifier
        interrupcion_inicio: Start time of the interruption from REPORTE DINAMICO 335 
        interrupcion_fin: End time of the interruption from REPORTE DINAMICO 335 
    
    Returns:
        Total clock stop by tickets
    
    """   
    df_sga_paradas['nro_incidencia'] = df_sga_paradas['nro_incidencia'].astype(str)
    nro_incidencia_stops = df_sga_paradas[df_sga_paradas['nro_incidencia'] == nro_incidencia].copy()

    if nro_incidencia_stops.empty:
        logger.info(f"No clock stops found for incident {nro_incidencia}")
        return 0.0
    
    clock_stops = []

    for _, stop in nro_incidencia_stops.iterrows():
        start_date = stop.get('startdate')
        end_date = stop.get('enddate')

        if pd.isna(start_date):
            logger.warning(f"Skipping record with missing start date for incident {nro_incidencia}")
            continue

        if start_date < interruption_start:
            logger.info(f"Adjusting start time to interruption en for incident {nro_incidencia}")
            start_date = interruption_start

        if not pd.isna(end_date):
            if end_date > interruption_end:
                logger.info(f"Adjusting end time to interruption en for incident {nro_incidencia}")
                end_date = interruption_end

            if start_date < end_date:
                clock_stops.append({
                    'start': start_date,
                    'end': end_date,
                    'nro_incidencia': nro_incidencia
                })
        else:
            clock_stops.append({
                'start': start_date,
                'end': end_date,
                'nro_incidencia': nro_incidencia
            })
    resolved_stops = resolve_clock_stop_overlaps(clock_stops)

    return resolved_stops




@log_exceptions
def has_repetition(text: str) -> bool:
    """
    Returns True if text contains o forbiden repetition of:
    - "Inmediatamente" twice, or
    - "A través" twice.
    """
    if not isinstance(text, str):
        return False
    
    patterns = [
        r'(?i)\b(inmediatamente)\b.*\b\1\b',
        r'(?i)\b(a través)\b.*\b\1\b',
    ]
    return any(re.search(p, text) for p in patterns )
   



# EXTRACT RANGE DATOS FECHA INICIO Y FIN FROM INDISPONIBILIDAD EXCEL Y SGA 

@log_exceptions
def extract_date_range_last(text: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Devuelve (fecha_hora_inicio, fecha_hora_fin) extraídas
    de las dos últimas líneas que contengan un
    'DD/MM/YYYY hh:mm' (con o sin 'a las'/'horas').
    Si falla el parseo, retorna (None, None).
    """
    if not isinstance(text, str):
        return (None, None)

    normalized = text.replace('\r', ' ')

    lines = [ln.strip() for ln in normalized.splitlines() if ln.strip()]
    if len(lines) < 2:
        return (None, None)

    dt_rx = re.compile(
        r'(\d{1,2}/\d{1,2}/\d{4})'    # fecha
        r'\s*(?:a las\s*)?'           # opcional “a las”
        r'(\d{1,2}:\d{2})'            # hora
        r'(?:\s*horas\.?)?',          # opcional “horas”
        re.IGNORECASE
    )


    date_lines = [ln for ln in lines if dt_rx.search(ln)]
    if len(date_lines) < 2:
        return (None, None)

    start_line, end_line = date_lines[-2], date_lines[-1]

    def parse(line: str) -> Optional[str]:
        m = dt_rx.search(line)
        return f"{m.group(1)} {m.group(2)}" if m else None

    inicio = parse(start_line)
    fin    = parse(end_line)
    return (inicio, fin) if inicio and fin else (None, None)



@log_exceptions
def extract_date_range_body(text: str) -> Tuple[Optional[str], Optional[str]]:
    if not isinstance(text, str):
        return (None, None)

  
    lines = text.splitlines()
    meta_pat = re.compile(r'(?i)^fecha y hora(?: de)? (?:inicio|fin)\s*:')
   
    while lines and (not lines[-1].strip() or meta_pat.match(lines[-1].strip())):
        lines.pop()
    clean_body = "\n".join(lines)

   
    narr_pat = re.compile(
        r'el día\s*(\d{2}/\d{2}/\d{4})\s*a las\s*(\d{2}:\d{2})',
        flags=re.IGNORECASE
    )

    combined_pat = re.compile(
        r'el(?:\s*d[ií]a)?\s*'           # “el” + opcional “ día”
        r'(\d{2}/\d{2}/\d{4})\s*'        # fecha
        r'(?:a las\s*)?'                 # opcional “a las”
        r'(\d{1,2}:\d{2})'               # hora
        r'(?:\s*horas?)?',               # opcional “ horas”
        flags=re.IGNORECASE
    )

    matches = combined_pat.findall(clean_body)

    if not matches:
        return (None, None)

    def _normalize(fecha: str, hora: str) -> str:
        d, m, y = fecha.split('/')
        h, mm = hora.split(':')
        d = d.zfill(2)
        m = m.zfill(2)
        h = h.zfill(2)
        mm = mm.zfill(2)
        return f"{d}/{m}/{y} {h}:{mm}"


    if matches:
        start_date, start_time = matches[0]
        end_date,   end_time   = matches[-1]
        return (_normalize(start_date, start_time), _normalize(end_date, end_time))

   


# WORD REPORTE TECNICO
@log_exceptions
def extract_tecnico_reports(path_docx: str) -> pd.DataFrame:
    """
    Extrae cada sección 'REPORTE TÉCNICO Nº X' en una fila.
    Saca CUISMP, Tipo Caso, Observación, Determinación de la causa, 
    Medidas tomadas, y además PARSEA la sección de MEDIDAS para 
    capturar Fecha y hora inicio / Fecha y hora fin que aparecen justo
    debajo de ese título.
    """
    doc = Document(path_docx)
    paragraphs = [p.text.strip() for p in doc.paragraphs]

    
    heading_pat = re.compile(r"REPORTE T[EÉ]CNICO Nº\s*(\d+)", re.IGNORECASE)
    headings: List[Dict] = []
    for idx, text in enumerate(paragraphs):
        m = heading_pat.search(text)
        if m:
            headings.append({"index": idx, "nro_incidencia": m.group(1)})
    if not headings:
        raise ValueError("No se encontró ‘REPORTE TÉCNICO Nº ...’ en el documento")


    headings.append({"index": len(paragraphs), "nro_incidencia": None})

    records: List[Dict] = []

    common_patterns = {
        "CUISMP":                   re.compile(r"CUISMP\s*(?:\:|\s)\s*(\d+)", re.IGNORECASE),
        "Tipo Caso":                re.compile(r"Tipo Caso\s*:\s*(.+)",    re.IGNORECASE),
        "Observación":              re.compile(r"Observaci.o.n\s*:\s*(.+)", re.IGNORECASE),
        "DETERMINACIÓN DE LA CAUSA":re.compile(r"Determinaci.o.n de la causa\s*:\s*(.+)", re.IGNORECASE),  
    }

    medidas_title_pat = re.compile(r"^MEDIDAS CORRECTIVAS Y/O PREVENTIVAS TOMADAS", re.IGNORECASE)

    inicio_pat = re.compile(r"^Fecha y hora inicio\s*:\s*(.+)$", re.IGNORECASE)
    fin_pat    = re.compile(r"^Fecha y hora fin\s*:\s*(.+)$",    re.IGNORECASE)


    for head, nxt in zip(headings, headings[1:]):
        start, end = head["index"], nxt["index"]
        block_lines = paragraphs[start:end]
        block_text  = "\n".join(block_lines)

 
        row: Dict[str, str] = {
            "nro_incidencia": head["nro_incidencia"],
            "CUISMP":         "",
            "Tipo Caso":      "",
            "Observación":    "",
            "DETERMINACION DE LA CAUSA": "",
            "MEDIDAS CORRECTIVAS Y/O PREVENTIVAS TOMADAS": "" ,  # texto completo si quieres
            "Fecha y hora inicio": "",
            "Fecha y hora fin":    "",
        }

        for key, pat in common_patterns.items():
            m = pat.search(block_text)
            if m:
                row[key] = m.group(1).strip()

        med_idx = None
        for i, line in enumerate(block_lines):
            if medidas_title_pat.match(line):
                med_idx = i
                break
        if med_idx is not None:

            meds = []
            for sub in block_lines[med_idx+1:]:
                if not sub:
                   
                    break
                meds.append(sub)
              
                mi = inicio_pat.match(sub)
                if mi:
                    row["Fecha y hora inicio"] = mi.group(1).strip()
                mf = fin_pat.match(sub)
                if mf:
                    row["Fecha y hora fin"] = mf.group(1).strip()
            row["MEDIDAS CORRECTIVAS Y/O PREVENTIVAS TOMADAS"] = " ".join(meds).strip()

        records.append(row)


    df = pd.DataFrame.from_records(records)

    return df



@log_exceptions
def extract_tecnico_reports_without_hours_last_dates(path_docx: str) -> pd.DataFrame:
    """
    Extrae cada sección 'REPORTE TÉCNICO Nº X' en una fila.
    Saca CUISMP, Tipo Caso, Observación, Determinación de la causa, 
    Medidas tomadas, y además PARSEA la sección de MEDIDAS para 
    capturar Fecha y hora inicio / Fecha y hora fin que aparecen justo
    debajo de ese título.
    """
    doc = Document(path_docx)
    paragraphs = [p.text.strip() for p in doc.paragraphs]

    
    heading_pat = re.compile(r"REPORTE T[EÉ]CNICO Nº\s*(\d+)", re.IGNORECASE)
    headings: List[Dict] = []
    for idx, text in enumerate(paragraphs):
        m = heading_pat.search(text)
        if m:
            headings.append({"index": idx, "nro_incidencia": m.group(1)})
    if not headings:
        raise ValueError("No se encontró ‘REPORTE TÉCNICO Nº ...’ en el documento")


    headings.append({"index": len(paragraphs), "nro_incidencia": None})

    records: List[Dict] = []

    common_patterns = {
        "CUISMP":                   re.compile(r"CUISMP\s*(?:\:|\s)\s*(\d+)", re.IGNORECASE),
        "Tipo Caso":                re.compile(r"Tipo Caso\s*:\s*(.+)",    re.IGNORECASE),
        "Observación": re.compile(r"Observaci[oó]n\s*:\s*(.+)", re.IGNORECASE),

    }

    medidas_title_pat = re.compile(r"^MEDIDAS CORRECTIVAS Y/O PREVENTIVAS TOMADAS", re.IGNORECASE)


    pattern = (
        r'^(?:Fecha y Hora|Hora)(?:\s+de)?\s+'    # “Fecha y Hora” o “Hora”, con “ de” opcional
        r'(Inicio|Fin):\s*'                      # “Inicio:” o “Fin:”
        r'(\d{1,2}/\d{1,2}/\d{4})'                # grupo(2) = fecha
        r'\s*(?:a las\s*)?'                      # “a las” opcional
        r'(\d{1,2}:\d{2})'                       # grupo(3) = hora
        r'(?:\s+horas\.?)?'                      # “horas” opcional
    )
    rx = re.compile(pattern, re.IGNORECASE)


    for head, nxt in zip(headings, headings[1:]):
        start, end = head["index"], nxt["index"]
        block_lines = paragraphs[start:end]
        block_text  = "\n".join(block_lines)

 
        row: Dict[str, str] = {
            "nro_incidencia": head["nro_incidencia"],
            "CUISMP":         "",
            "Tipo Caso":      "",
            "Observación":    "",
            "DETERMINACIÓN DE LA CAUSA": "",
            "MEDIDAS CORRECTIVAS Y/O PREVENTIVAS TOMADAS": "" ,  # texto completo si quieres
            "Fecha y hora inicio": "",
            "Fecha y hora fin":    "",
        }

        cause_idx = next(
        (i for i, line in enumerate(block_lines)
         if re.match(r"Determinaci[oó]n\s+de\s+la\s+causa\s*:?\s*$",
                     line, re.IGNORECASE)),
        None
        )
        if cause_idx is not None:
            for j in range(cause_idx + 1, len(block_lines)):
                text = block_lines[j].strip()
                if text:
                    row["DETERMINACIÓN DE LA CAUSA"] = text
                    break    


        for key, pat in common_patterns.items():
            m = pat.search(block_text)
            if m:
                row[key] = m.group(1).strip()

        med_idx = None
        for i, line in enumerate(block_lines):
            if medidas_title_pat.match(line):
                med_idx = i
                break
        if med_idx is not None:

            meds = []
            for sub in block_lines[med_idx+1:]:
                if not sub:
                   
                    #break
                    continue
              
                mi = rx.match(sub)
                if mi:
                    if mi.group(1).lower() == "inicio":
                        row["Fecha y hora inicio"] = f"{mi.group(2)} {mi.group(3)}"
                    elif mi.group(1).lower() == "fin":
                        row["Fecha y hora fin"] = f"{mi.group(2)} {mi.group(3)}"
                    continue

                meds.append(sub)

            row["MEDIDAS CORRECTIVAS Y/O PREVENTIVAS TOMADAS"] = " ".join(meds).strip()

        records.append(row)


    df = pd.DataFrame.from_records(records)

    return df


# WORD ANEXOS
@log_exceptions
def extract_indisponibilidad_anexos(path_docx: str) ->  Optional[pd.DataFrame]:
    """
    Extrae de un .docx que contiene varios ANEXO X – TICKET Y:
      - la línea de introducción ("Se tuvo indisponibilidad...")
      - las líneas de periodos ("DD/MM/YYYY hh:mm hasta el día DD/MM/YYYY hh:mm")
      - la línea del total de horas (“(Total de horas sin acceso… HH:MM horas)”)
    Retorna un DataFrame con columnas:
      ['ticket', 'indisponibilidad_header', 'indisponibilidad_periodos', 'indisponibilidad_total']
    """

    doc = Document(path_docx)


    raw = [p.text for p in doc.paragraphs]
    paragraphs_list = []
    for para in raw:
        for line in para.splitlines():
            text = line.strip()
            if text:
                paragraphs_list.append(text)

    anexo_pattern_ticket = re.compile(r"ANEXO\s+\d+\s+–\s+TICKET\s+(\d+)", re.IGNORECASE)
    linea0_pat_pattern   = re.compile(r"^Se tuvo indisponibilidad por parte del cliente.*", re.IGNORECASE)
    periodo_pattern      = re.compile(
        r"^\d{1,2}/\d{1,2}/\d{4}\s+\d{1,2}:\d{2}.*hasta el día\s+\d{1,2}/\d{1,2}/\d{4}\s+\d{1,2}:\d{2}.*$",
        re.IGNORECASE
    )
    total_hour_pattern   = re.compile(r"^\(Total de horas sin acceso a la sede:\s*(\d{1,3}:\d{2})\s*horas\)", re.IGNORECASE)

    
    _day_pad_start = re.compile(r"^(\d)(?=/)")
    _day_pad_end   = re.compile(r"(?<=hasta el día\s)(\d)(?=/)", re.IGNORECASE)
    _hour_pad      = re.compile(r"\b(\d)(?=:\d{2}(?::\d{2})?)")

    def pad_periodo(line: str) -> str:
        
        line = _day_pad_start.sub(lambda m: m.group(1).zfill(2), line)
        line = _day_pad_end.sub(lambda m: m.group(1).zfill(2), line)
        return _hour_pad.sub(lambda m: m.group(1).zfill(2), line)
    
    def pad_total(tm: str) -> str:
        return _hour_pad.sub(lambda m: m.group(1).zfill(2), tm)

    records: List[Dict] = []

    for idx, text in enumerate(paragraphs_list):
        m_ticket = anexo_pattern_ticket.search(text)
        if not m_ticket:
            continue
        ticket = m_ticket.group(1)

        linea0  = ""
        periodos: List[str] = []
        footer   = ""
        total = ""


        for line in paragraphs_list[idx + 1:]:
            if anexo_pattern_ticket.search(line):
                break

            if linea0_pat_pattern.match(line):
                linea0 = line

            if periodo_pattern.match(line):
                periodos.append(pad_periodo(line))

            if total_hour_pattern.search(line):
                footer = line
                
            m_total =  total_hour_pattern.search(line)
            if m_total:
                total = m_total.group(1)
                break
                

        records.append({
            "ticket": ticket,
            "indisponibilidad_header": linea0,
            "indisponibilidad_periodos": "\n".join(periodos),
            "indisponibilidad_footer": footer,
            "indisponibilidad_total": total
        })
    
    if not records:
        return None

    return pd.DataFrame.from_records(records)




























# @log_exceptions
# def is_langtool_clean(text:str) -> bool:

#     """
#     Return True if LanguageTool reports zero issues in the text.
#     """
#     if not isinstance(text, str) or not text.strip():
#         return True
    
#     matches = _tool_es.check(text)
#     return len(matches) == 0

