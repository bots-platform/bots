# word_extractor.py

import pandas as pd
from docx import Document
from typing import List, Dict
from utils.logger_config import get_sga_logger

logger = get_sga_logger()

def log_exceptions(func):
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            logger.error(f"Error in {func.__name__}: {e}", exc_info=True)
            raise
    return wrapper

@log_exceptions
def extract_averias_table(path_docx: str) -> pd.DataFrame:
    """
    Load the specified table from a .docx and return it as a DataFrame.
    index = 4 (CLARO)
    index = 5 (CLIENTE)
    index =  (TERCEROS)
    We assume the first row is the header.
    """

    table_index_list = [4,5,6]

    doc = Document(path_docx)
    all_dfs: List[pd.DataFrame] = []
   

    for table_index in table_index_list:

        if table_index >= len(doc.tables):
            raise IndexError(f"Document only has {len(doc.tables)} tables.")
        
        table = doc.tables[table_index]

        headers = [cell.text.strip() for cell in table.rows[0].cells]

        records: List[Dict] = []
        for row in table.rows[1:]:
            vals = [cell.text.strip() for cell in row.cells]

            if len(vals) < len(headers):
                vals += [""] * (len(headers) - len(vals))
            records.append(dict(zip(headers, vals)))

        df_part = pd.DataFrame.from_records(records)
        
        if table_index == 4:
            df_part['responsable'] = 'CLARO'
        elif table_index == 5:
            df_part['responsable'] = 'CLIENTE'
        else: 
            df_part['responsable'] = 'TERCEROS'

        all_dfs.append(df_part)

    result = pd.concat(all_dfs, ignore_index=True)

    return result
